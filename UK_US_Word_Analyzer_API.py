from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.background import BackgroundTasks
from typing import Dict, List, Tuple, Any
import docx2txt
from docx import Document
from collections import defaultdict
import os
import json
import tempfile
import re
from fastapi.middleware.cors import CORSMiddleware
app = FastAPI(title="Document Analyzer API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Or specify allowed origins
    allow_credentials=True,
    allow_methods=["*"],  # Or ["GET", "POST"]
    allow_headers=["*"],
)
def load_dictionary(dict_path: str) -> set:
    """Load dictionary words from a file"""
    with open(dict_path, 'r', encoding='utf-8') as f:
        return set(word.strip().lower() for word in f.readlines() if word.strip())
async def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from DOCX file bytes
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        HTTPException: If file processing fails
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name
    
    try:
        return docx2txt.process(temp_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX processing failed: {str(e)}")
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)

async def analyze_docx(file_content: bytes, us_words: set, uk_words: set) -> Dict[str, Any]:
    """
    Analyze DOCX content for UK/US word matches
    """
    
    try:
        # Process DOCX file
        text = await extract_text_from_docx(file_content)
        words = text.lower().split()
        
        # Initialize results
        results = {
            'us': {'total': 0, 'words': defaultdict(int)},
            'uk': {'total': 0, 'words': defaultdict(int)},
            'total_words_in_document': len(words)
        }
        
        # Check each word against dictionaries
        for word in words:
            clean_word = word.strip(".,!?()\"'")
            if clean_word:
                if clean_word in us_words:
                    results['us']['words'][clean_word] += 1
                    results['us']['total'] += 1
                if clean_word in uk_words:
                    results['uk']['words'][clean_word] += 1
                    results['uk']['total'] += 1
        
        # Convert defaultdict to regular dict
        results['us']['words'] = dict(results['us']['words'])
        results['uk']['words'] = dict(results['uk']['words'])
        
        return results
    
    except HTTPException:
        raise  # Re-raise existing HTTP exceptions
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

def generate_report_content(results: dict) -> str:
    """Generate report content as string"""
    report_content = "=== DOCUMENT ANALYSIS REPORT ===\n"
    report_content += f"Total Words: {results['total_words_in_document']}\n\n"
    
    report_content += "US Words \n"
    report_content += f"Total Words : {results['us']['total']}\n"
    for word, count in results['us']['words'].items():
        report_content += f"{word}: {count}\n"
    
    report_content += "\nUK Words \n"
    report_content += f"Total Words : {results['uk']['total']}\n"
    for word, count in results['uk']['words'].items():
        report_content += f"{word}: {count}\n"
    
    return report_content

def cleanup_tempfile(path: str):
    """Clean up temporary file"""
    if os.path.exists(path):
        os.unlink(path)

@app.post("/analyze", response_model=Dict[str, Any])
async def analyze_document(
    file: UploadFile = File(...) 
):
    """
    Analyze a DOCX file for UK/US word usage
    """
    try:
        
        # Check file type
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only DOCX files are supported")
        us_dict_path = os.getenv("US_DICT_PATH", "us_dict.txt")
        uk_dict_path = os.getenv("UK_DICT_PATH", "uk_dict.txt")
        # Load dictionaries
        if not os.path.exists(us_dict_path):
            raise HTTPException(status_code=404, detail="US dictionary file not found")
        if not os.path.exists(uk_dict_path):
            raise HTTPException(status_code=404, detail="UK dictionary file not found")
            
        us_words = load_dictionary(us_dict_path)
        uk_words = load_dictionary(uk_dict_path)
        
        # Analyze document
        file_content = await file.read()
        results = await analyze_docx(file_content, us_words, uk_words)
        
        return JSONResponse(content=results)
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-and-download")
async def analyze_and_download(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    """
    Analyze a DOCX file and return a text report download
    """
    try:
        # Check file type
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only DOCX files are supported")
        
        # Load dictionaries
        us_dict_path = os.getenv("US_DICT_PATH", "us_dict.txt")
        uk_dict_path = os.getenv("UK_DICT_PATH", "uk_dict.txt")
        if not os.path.exists(us_dict_path):
            raise HTTPException(status_code=404, detail="US dictionary file not found")
        if not os.path.exists(uk_dict_path):
            raise HTTPException(status_code=404, detail="UK dictionary file not found")
            
        us_words = load_dictionary(us_dict_path)
        uk_words = load_dictionary(uk_dict_path)
        
        # Analyze document
        file_content = await file.read()
        results = await analyze_docx(file_content, us_words, uk_words)
        
        # Generate report content
        report_content = generate_report_content(results)
        
        # Create temporary report file
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as tmp:
            tmp.write(report_content)
            tmp_path = tmp.name
        
        # Schedule cleanup after response is sent
        background_tasks.add_task(cleanup_tempfile, tmp_path)
        
        return FileResponse(
            tmp_path,
            media_type="text/plain",
            filename="word_analysis_report.txt",
            headers={"Content-Disposition": "attachment; filename=word_analysis_report.txt"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@app.post("/abbreviation")
async def abbreviation(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only DOCX files are supported")
        file_content = await file.read()
        text = await extract_text_from_docx(file_content)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            # Write the uploaded file's content to the temp file
            temp_file.write(file_content)
            temp_path = temp_file.name
        doc = Document(temp_path)    
        #extract_abbreviation_list(text)
        responseValue = await extract_abbreviation_list(text)
        serial_commas = find_serial_commas(doc)
        responseValue['serial_commas'] = serial_commas['serial_commas']
        return JSONResponse(content=responseValue)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/abbreviation-and-download")
async def abbreviation_and_download(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only DOCX files are supported")
        file_content = await file.read()
        text = await extract_text_from_docx(file_content)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            # Write the uploaded file's content to the temp file
            temp_file.write(file_content)
            temp_path = temp_file.name
        doc = Document(temp_path)    
        #extract_abbreviation_list(text)
        responseValue = await extract_abbreviation_list(text)
        serial_commas = find_serial_commas(doc)
        responseValue['serial_commas'] = serial_commas['serial_commas']
        # Create temporary report file
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as tmp:
            json.dump(responseValue,tmp, indent=4)
            tmp_path = tmp.name
        
        # Schedule cleanup after response is sent
        background_tasks.add_task(cleanup_tempfile, tmp_path)
        
        return FileResponse(
            tmp_path,
            media_type="text/plain",
            filename="word_abbreviation_report.txt",
            headers={"Content-Disposition": "attachment; filename=word_abbreviation_report.txt"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def extract_abbreviation_list(document_content):
    try:
        abbreviation_list = []
        abbreviation_count = []
        all_found_abbreviations = set()

        pattern = r"(([A-Z]+ [0-9]+)|\b([A-Z]+(-?[A-Z0-9]+)*(s)?)\b)|([a-z]*[A-Z][a-z]*)+"
        matches = re.findall(pattern, document_content)

        for match_group in matches:
            found_abbreviation = next((m for m in match_group if m), "").strip()
            found_abbreviation = re.sub(r"[^-a-zA-Z0-9 ]", "", found_abbreviation)
            found_abbreviation = re.sub(r"(^[- ]|[- ]$)", "", found_abbreviation)

            if found_abbreviation.endswith("s"):
                found_abbreviation = found_abbreviation[:-1]

            if (2 <= len(found_abbreviation) <= 6 and
                    len(re.findall(r"[A-Z]", found_abbreviation)) > 1 and
                    found_abbreviation not in all_found_abbreviations):

                count_pattern = rf"\b{re.escape(found_abbreviation)}s?\b"
                count = len(re.findall(count_pattern, document_content))

                abbreviation_list.append(found_abbreviation)
                abbreviation_count.append(count)
                all_found_abbreviations.add(found_abbreviation)

        expansion_array = resolve_search_priority(abbreviation_list, document_content)

        # Assemble the desired response format
        result = []
        for idx, abbr in enumerate(abbreviation_list):
            expansion = expansion_array[idx] if idx < len(expansion_array) else ""
            expansion = expansion.lstrip('|').split('|')[0] if expansion else ""
            result.append({
                "abbreviation": abbr,
                "full_form": expansion,
                "occurrences": abbreviation_count[idx]
            })

        return {"abbreviations_found": result}
    except Exception as e:
        print(f"Error resolving abbreviation expansions in extract_abbreviation_list def: {e}")
def resolve_search_priority(abbreviation_list, document_full_content):
    expansion_array = [''] * (len(abbreviation_list) + 1)
    #db_expansion_array = [''] * (len(abbreviation_list) + 1)

    try:
        expansion_array = get_expansion_from_document(abbreviation_list, document_full_content)

    except Exception as e:
        print(f"Error resolving abbreviation expansions in resolve_search_priority def: {e}")

    return expansion_array
def get_expansion_from_document(abbreviation_list, document_full_content):
    """
    Retrieves possible expansions for abbreviations from the document content.

    :param abbreviation_list: List of abbreviations (e.g. ['NASA', 'WHO']).
    :param document_full_content: The full document text.
    :return: List of expansions (same length as abbreviation_list + 1).
    """
    try:
        # Initialize expansion list (+1 for compatibility with original C# structure)
        expansion_array = [''] * (len(abbreviation_list) + 1)

        # Retrieve expansions using a custom heuristic
        expansion_array = new_approach_retrieve_expansion(abbreviation_list, document_full_content, expansion_array)

        # Add XML tags or formatting
        #expansion_array = add_xml_tags(expansion_array, document_full_content)

        return expansion_array

    except Exception as e:
        print(f"Error in get_expansion_from_document in: {e}")
        return [''] * (len(abbreviation_list) + 1)
def new_approach_retrieve_expansion(abbreviation_list, document_text, expansion_array=None):
    try:
        if expansion_array is None:
            expansion_array = [''] * (len(abbreviation_list) + 1)

        preposition_list = 'at|by|for|of|in|on|to|and|the|at'
        list_of_all_expansions = set()

        for idx in range(0, len(abbreviation_list)):
            abbr = abbreviation_list[idx]
            filtered_text = ''
            is_expansion = False

            if len(abbr) > 2:
                # Try abbreviation as-is (e.g., "World Health Organization (WHO)")
                filtered_text = new_approach_expansion_as_is_form(document_text, abbr, preposition_list)
                filtered_text = clean_surrounding_nonword(filtered_text)

                if filtered_text and filtered_text not in list_of_all_expansions:
                    list_of_all_expansions.add(filtered_text)
                    expansion_array[idx] += f"|{filtered_text.strip()}"
                else:
                    pattern = search_pattern_1(abbr, preposition_list) if re.search(r'[a-z]', abbr) else search_pattern(abbr, preposition_list)
                    matches = re.findall(pattern, document_text, flags=re.IGNORECASE)
                    for match_tuple in matches:
                        for match in match_tuple:
                            if not isinstance(match, str):
                                continue
                            filtered_text = clean_surrounding_nonword(match)
                            if not filtered_text or filtered_text in list_of_all_expansions:
                                continue

                            s_text = filtered_text
                            if s_text.lower().startswith(abbr.lower()):
                                s_text = s_text[len(abbr):].strip()

                            if len(s_text) > len(abbr):
                                i_pos = 0
                                is_expansion = True
                                for c in abbr:
                                    found = re.search(re.escape(c), s_text[i_pos:], flags=re.IGNORECASE)
                                    if found:
                                        i_pos += found.start() + 1
                                    else:
                                        is_expansion = False
                                        break

                                if is_expansion:
                                    list_of_all_expansions.add(filtered_text)
                                    expansion_array[idx] += f"|{filtered_text.strip()}"
                                    break
            elif len(abbr) == 2:
                filtered_text = new_approach_expansion_as_is_form(document_text, abbr, preposition_list)
                filtered_text = clean_surrounding_nonword(filtered_text)

                if filtered_text and filtered_text not in list_of_all_expansions:
                    s_text = filtered_text
                    if s_text.lower().startswith(abbr.lower()):
                        s_text = s_text[len(abbr):].strip()

                    if len(s_text) > len(abbr):
                        i_pos = 0
                        is_expansion = True
                        for c in abbr:
                            found = re.search(re.escape(c), s_text[i_pos:], flags=re.IGNORECASE)
                            if found:
                                i_pos += found.start() + 1
                            else:
                                is_expansion = False
                                break

                        if is_expansion:
                            list_of_all_expansions.add(filtered_text)
                            expansion_array[idx] += f"|{filtered_text.strip()}"
                        else:
                            expansion_array[idx] = ""
                    else:
                        expansion_array[idx] = ""
                else:
                    expansion_array[idx] = ""

        filter_searched_expansion(expansion_array, abbreviation_list, document_text)
        return expansion_array
    except Exception as e:
        print(f"Error in new_approach_retrieve_expansion def: {e}")
def filter_searched_expansion(expansion_array: list[str], abbreviation_list: list[str], document_full_content: str) -> None:
    try:
        # 1. Filter based on casing
        for i in range(1, len(expansion_array)):
            if expansion_array[i]:
                split_expansion = expansion_array[i].split('|')
                filtered_text = ""
                for item in split_expansion[1:]:
                    temp_string = re.sub(r"[a-z\W]", "", item)
                    if temp_string == abbreviation_list[i]:
                        filtered_text += "|" + item
                if filtered_text:
                    expansion_array[i] = filtered_text

        # 2. Filter based on same paragraph presence
        split_para = document_full_content.split('\n')
        for i in range(1, len(expansion_array)):
            if expansion_array[i]:
                split_expansion = expansion_array[i].split('|')
                if len(split_expansion) > 2:
                    filtered_text = ""
                    for item in split_expansion[1:]:
                        for para in split_para:
                            if (re.search(r"\W" + re.escape(item) + r"\W", para, re.IGNORECASE) and
                                re.search(r"\W" + re.escape(abbreviation_list[i]) + r"\W", para)):
                                filtered_text += "|" + item
                                break  # Stop searching after the first matching paragraph
                    expansion_array[i] = filtered_text if filtered_text else ""
    except Exception as e:
        print(f"Error resolving abbreviation expansions in filter_searched_expansion def: {e}")

def new_approach_expansion_as_is_form(previous_text, abbreviation, preposition_list):
    try:
        filtered_text = ''
        
        # Build the pattern for expansion
        pattern = build_single_word_regex_pattern(abbreviation, preposition_list)

        # Try pattern like: expansion text + abbreviation (e.g., World Health Organization (WHO))
        pattern_temp = rf"{pattern}(\W)+{re.escape(abbreviation)}(s)?(\W)+"
        match = re.search(pattern_temp, previous_text, re.IGNORECASE)
        
        if match:
            filtered_text = match.group(0)
            filtered_text = re.sub(rf"\({re.escape(abbreviation)}\)", "", filtered_text, flags=re.IGNORECASE)
        else:
            # Try pattern like: abbreviation + expansion (e.g., WHO (World Health Organization))
            pattern_temp = rf"{re.escape(abbreviation)}(\W)*{pattern}(\W)+"
            match = re.search(pattern_temp, previous_text, re.IGNORECASE)
            if match:
                filtered_text = match.group(0)
                filtered_text = re.sub(re.escape(abbreviation), "", filtered_text, flags=re.IGNORECASE)
        
        if not filtered_text:
            # Try pattern like: expansion , abbreviation
            pattern_temp = rf"{pattern}(\W)*,{re.escape(abbreviation)}"
            match = re.search(pattern_temp, previous_text, re.IGNORECASE)
            if match:
                filtered_text = match.group(0)
        if not filtered_text:
            pattern_temp = rf"{pattern}(\W)+({preposition_list})?(\s)?[A-Z][a-z\-\'\’]+(\s)?(\W){re.escape(abbreviation)}(\s)?(\W)+"
            match = re.search(pattern_temp, previous_text, re.IGNORECASE)
            if match:
                filtered_text = match.group(0)
                #filtered_text = re.sub(rf"(\W)?{re.escape(abbreviation)}(s)?(\W)?", "", filtered_text, flags=re.IGNORECASE)
                filtered_text = check_abbreviation_match(abbreviation, filtered_text)

        # Clean abbreviation from result
        if filtered_text:
            filtered_text = re.sub(rf"(\W)?{re.escape(abbreviation)}(s)?(\W)?", "", filtered_text, flags=re.IGNORECASE)

        return filtered_text.strip()
    except Exception as e:
        print(f"Error resolving abbreviation expansions in new_approach_expansion_as_is_form def: {e}")
def check_abbreviation_match(abbreviation, filtered_text):
    # List of words to ignore when checking first letters
    ignore_words = {'at', 'by', 'for', 'of', 'in', 'on', 'to', 'and', 'the'}
    
    # Remove the abbreviation in parentheses if present
    filtered_text = re.sub(rf"\({re.escape(abbreviation)}\)", "", filtered_text, flags=re.IGNORECASE)
    filtered_text = re.sub(r'[^a-zA-Z\s]', ' ', filtered_text)
    filtered_text = re.sub(rf"\b{re.escape(abbreviation)}\b", "", filtered_text, flags=re.IGNORECASE).strip()
    if re.search(rf'\b{re.escape(abbreviation)}\b', filtered_text, flags=re.IGNORECASE):
        return ""
    # Split the filtered text into words
    words = filtered_text.split()
    
    # Filter out ignored words and get first letters
    first_letters = []
    for word in words:
        lower_word = word.lower()
        if lower_word not in ignore_words:
            if word:  # check if word is not empty
                first_letters.append(word[0].lower())
    
    # Check if each character in abbreviation matches the first letters
    if len(abbreviation) > len(first_letters):
        return ""
    
    for i in range(len(abbreviation)):
        if i >= len(first_letters):
            return ""
        if abbreviation[i].lower() != first_letters[i]:
            return ""
    
    return filtered_text        
def build_single_word_regex_pattern(abbreviation: str, preposition_list: str) -> str:
    try:
        abbr_chars = list(abbreviation)
        pattern = (
            f"[{abbr_chars[0].upper()}{abbr_chars[0].lower()}][a-zA-Z\\-]+(\\s)?"
            f"({preposition_list})?(\\s)?"
        )

        for i in range(1, len(abbr_chars)):
            ch = abbr_chars[i]

            if i < len(abbr_chars) - 1:
                if re.match(r"[0-9]", ch):
                    pattern += (
                        f"([{ch.upper()}][a-z\\-\\'\\’]*)(\\s)?"
                        f"({preposition_list})?(\\s)?"
                    )
                elif re.match(r"[- ]", ch):
                    pattern += (
                        f"([{ch.upper()}][a-z\\-\\'\\’]*)?(\\s)?"
                        f"({preposition_list})?(\\s)?"
                    )
                else:
                    pattern += (
                        f"[{ch.upper()}{ch.lower()}][a-z\\-\\'\\’]+(\\s)?"
                        f"({preposition_list})?(\\s)?"
                    )
            else:
                if re.match(r"[- 0-9]", ch):
                    pattern += f"([{ch.upper()}][a-z\\-\\'\\’]*)"
                else:
                    pattern += f"[{ch.upper()}{ch.lower()}][a-z\\-\\'\\’]+"

        pattern = "(\\W)" + pattern
        return pattern
    except Exception as e:
        print(f"Error resolving abbreviation expansions in build_single_word_regex_pattern def: {e}")
def clean_surrounding_nonword(text):
    try:
        if not text:
            return ""
        return re.sub(r"^\W+|\W+$", "", text)
    except Exception as e:
        print(f"Error in clean_surrounding_nonword in clean_surrounding_nonword def: {e}")

def search_pattern_1(abbreviation: str, preposition_list: str) -> str:
    try:
        abbr_chars = list(abbreviation)
        pattern = (
            f"[{abbr_chars[0].upper()}{abbr_chars[0].lower()}][a-z\\'\\’\\-]*"
            f"({preposition_list})?"
        )

        for i in range(1, len(abbr_chars)):
            ch = abbr_chars[i]

            if i < len(abbr_chars) - 1:
                if re.match(r"[0-9]", ch):
                    pattern += (
                        f"(\\s)?([{ch.upper()}][a-z\\'\\’\\-]*)"
                        f"({preposition_list})?(\\s)?"
                    )
                elif re.match(r"[- ]", ch):
                    pattern += (
                        f"(\\s)?([{ch.upper()}][a-z\\'\\’\\-]*)?"
                        f"({preposition_list})?(\\s)?"
                    )
                else:
                    pattern += (
                        f"(\\s)?[{ch.upper()}{ch.lower()}][a-z\\'\\’\\-]*"
                        f"({preposition_list})?(\\s)?"
                    )
            else:
                if re.match(r"[0-9]", ch):
                    pattern += (
                        f"(\\s)?([{ch.lower()}][a-zA-Z\\-\\'\\’]*)(\\W)"
                    )
                elif re.match(r"[- ]", ch):
                    pattern += (
                        f"(\\s)?([{ch.lower()}][a-zA-Z\\-\\'\\’]*)?(\\W)"
                    )
                else:
                    pattern += (
                        f"(\\s)?[{ch.upper()}{ch.lower()}][a-z\\-\\'\\’]*(\\W)"
                    )

        pattern = r"(\W)" + pattern
        return pattern
    except Exception as e:
        print(f"Error resolving abbreviation expansions in search_pattern_1 def: {e}")
def search_pattern(abbreviation: str, preposition_list: str) -> str:
    try:
        abbr_chars = list(abbreviation)
        
        # Start the pattern with the first character's word
        pattern = (
            f"[{abbr_chars[0].upper()}{abbr_chars[0].lower()}][a-zA-Z\\'\\’\\-]+"
            f"({preposition_list})?"
        )

        for i in range(1, len(abbr_chars)):
            ch = abbr_chars[i]

            if i < len(abbr_chars) - 1:
                if re.match(r"[0-9]", ch):
                    pattern += (
                        f"(\\s)?([{ch.upper()}][a-z\\'\\’\\-]*)"
                        f"({preposition_list})?(\\s)?"
                    )
                elif re.match(r"[- ]", ch):
                    pattern += (
                        f"(\\s)?([{ch.upper()}][a-z\\'\\’\\-]*)?"
                        f"({preposition_list})?(\\s)?"
                    )
                else:
                    pattern += (
                        f"(\\s)+[{ch.upper()}{ch.lower()}][a-z\\'\\’\\-]+"
                        f"({preposition_list})?(\\s)?"
                    )
            else:
                if re.match(r"[0-9]", ch):
                    pattern += (
                        f"(\\s)?([{ch.lower()}][a-zA-Z\\-\\'\\’]*)(\\W)"
                    )
                elif re.match(r"[- ]", ch):
                    pattern += (
                        f"(\\s)?([{ch.lower()}][a-zA-Z\\-\\'\\’]*)?(\\W)"
                    )
                else:
                    pattern += (
                        f"(\\s)+[{ch.upper()}{ch.lower()}][a-zA-Z\\-\\'\\’]+(\\W)"
                    )

        pattern = r"(\W)" + pattern
        return pattern
    except Exception as e:
        print(f"Error resolving abbreviation expansions in search_pattern def: {e}")
        
def find_serial_commas(doc) -> List[Dict[str, Any]]:
    try:
        """Find serial commas (Oxford commas) in the document.
        
        Args:
            doc (Document): The Word document object
            
        Returns:
            List[Dict[str, Any]]: List of dictionaries containing serial comma information
        """
        serial_commas = []
        serial_comma_pattern = r'\b[^,.]+,\s+[^,.]+,\s+(?:and|or)\s+[^,.]+\b'
        # Process paragraphs
        for para_index, paragraph in enumerate(doc.paragraphs, 1):
            matches = re.finditer(serial_comma_pattern, paragraph.text)
            for match in matches:
                serial_commas.append({
                    'text': match.group(),
                    'location': f'Paragraph {para_index}'
                })
        
        # Process tables
        for table_index, table in enumerate(doc.tables, 1):
            for row_index, row in enumerate(table.rows, 1):
                for cell_index, cell in enumerate(row.cells, 1):
                    matches = re.finditer(serial_comma_pattern, cell.text)
                    for match in matches:
                        serial_commas.append({
                            'text': match.group(),
                            'location': f'Table {table_index}, Row {row_index}, Cell {cell_index}'
                        })
        return {"serial_commas": serial_commas}
        #return serial_commas
    except Exception as e:
        print(f"Error resolving abbreviation expansions in search_pattern def: {e}")

@app.get("/health")
def health_check():
    return {"status": "ok"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
